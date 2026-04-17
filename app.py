"""
Skills Crafter - Minecraft Education WebSocket Dashboard.

A web application that connects to Minecraft Education Edition clients,
displays player activity logs, and shows a real-time 2D map of player positions.
"""

import io
import json
import os
import socket
import time
import uuid
from flask import Flask, render_template, request, jsonify, send_file
from flask_socketio import SocketIO
from minecraft_server import MinecraftWSServer
from settings_manager import (
    load_settings, save_settings, mask_api_key,
    load_rubrics, save_rubrics,
)

app = Flask(__name__)
app.config["SECRET_KEY"] = "skills-crafter-secret"
socketio = SocketIO(app, async_mode="threading", cors_allowed_origins="*")

mc_server = MinecraftWSServer(socketio=socketio)


def get_public_host():
    """Get the public IP for display in the connect string.
    Uses EXTERNAL_HOST env var. If it's a hostname, resolves it to an IP.
    Falls back to local IP for development."""
    ext = os.environ.get("EXTERNAL_HOST", "").strip()
    if ext:
        # If it's already an IP, return as-is
        try:
            socket.inet_aton(ext)
            return ext
        except socket.error:
            pass
        # It's a hostname — resolve to IP for Minecraft
        try:
            return socket.gethostbyname(ext)
        except socket.gaierror:
            return ext
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception:
        return "127.0.0.1"


def _get_llm_client():
    """Get an LLM client from the saved settings. Returns (client, provider, model)."""
    settings = load_settings()
    provider = settings.get("llm_provider", "")
    api_key = settings.get("llm_api_key", "")
    endpoint = settings.get("llm_endpoint", "")

    if not provider or not api_key:
        return None, None, None

    if provider == "openai":
        import openai
        return openai.OpenAI(api_key=api_key), "openai", "gpt-4o-mini"
    elif provider == "anthropic":
        import anthropic
        return anthropic.Anthropic(api_key=api_key), "anthropic", "claude-haiku-4-5-20251001"
    elif provider == "azure":
        import openai
        return openai.AzureOpenAI(
            api_key=api_key, azure_endpoint=endpoint, api_version="2024-02-01"
        ), "azure", "gpt-4o-mini"
    return None, None, None


def _llm_chat(prompt, system=None, max_tokens=2048, timeout=45):
    """Send a chat message to the configured LLM. Returns the response text."""
    import sys
    client, provider, model = _get_llm_client()
    if not client:
        raise ValueError("No LLM configured — add an API key in Settings")

    print(f"[LLM] Calling {provider}/{model} max_tokens={max_tokens}...", flush=True)
    t0 = time.time()

    try:
        if provider == "anthropic":
            kwargs = {"model": model, "max_tokens": max_tokens, "timeout": timeout,
                      "messages": [{"role": "user", "content": prompt}]}
            if system:
                kwargs["system"] = system
            resp = client.messages.create(**kwargs)
            text = resp.content[0].text
        else:
            messages = []
            if system:
                messages.append({"role": "system", "content": system})
            messages.append({"role": "user", "content": prompt})
            resp = client.chat.completions.create(model=model, messages=messages,
                                                   max_tokens=max_tokens, timeout=timeout)
            text = resp.choices[0].message.content

        elapsed = round(time.time() - t0, 1)
        print(f"[LLM] Response received in {elapsed}s ({len(text)} chars)", flush=True)
        return text

    except Exception as e:
        elapsed = round(time.time() - t0, 1)
        print(f"[LLM] FAILED after {elapsed}s: {e}", file=sys.stderr, flush=True)
        raise
        return resp.choices[0].message.content


# ─── Page Routes ───

@app.route("/")
def index():
    local_ip = get_public_host()
    settings = load_settings()
    return render_template(
        "index.html",
        local_ip=local_ip,
        ws_port=mc_server.port,
        server_running=mc_server.running,
        settings=settings,
        masked_key=mask_api_key(settings.get("llm_api_key", "")),
    )


# ─── Settings API ───

@app.route("/api/settings", methods=["GET"])
def get_settings():
    settings = load_settings()
    masked = dict(settings)
    masked["llm_api_key"] = mask_api_key(settings.get("llm_api_key", ""))
    return jsonify(masked)


@app.route("/api/settings", methods=["POST"])
def update_settings():
    data = request.get_json()
    current = load_settings()

    if "llm_provider" in data:
        current["llm_provider"] = data["llm_provider"]
    if "llm_endpoint" in data:
        current["llm_endpoint"] = data["llm_endpoint"]
    if "welcome_message" in data:
        current["welcome_message"] = data["welcome_message"]
    if "welcome_color" in data:
        current["welcome_color"] = data["welcome_color"]
    if "show_trace_paths" in data:
        current["show_trace_paths"] = data["show_trace_paths"]
    if "report_detail_level" in data:
        current["report_detail_level"] = max(1, min(5, int(data["report_detail_level"])))

    if "llm_api_key" in data and data["llm_api_key"] and "*" not in data["llm_api_key"]:
        current["llm_api_key"] = data["llm_api_key"]

    save_settings(current)
    mc_server.welcome_message = current["welcome_message"]
    mc_server.welcome_color = current["welcome_color"]

    masked = dict(current)
    masked["llm_api_key"] = mask_api_key(current.get("llm_api_key", ""))
    return jsonify({"status": "ok", "settings": masked})


@app.route("/api/settings/test-key", methods=["POST"])
def test_api_key():
    data = request.get_json()
    provider = data.get("provider", "")
    api_key = data.get("api_key", "")
    endpoint = data.get("endpoint", "")

    if not api_key:
        settings = load_settings()
        provider = provider or settings.get("llm_provider", "")
        api_key = settings.get("llm_api_key", "")
        endpoint = endpoint or settings.get("llm_endpoint", "")

    if not provider or not api_key:
        return jsonify({"connected": False, "error": "No provider or API key configured"})

    import time as _time
    max_retries = 3
    last_err = ""
    for attempt in range(max_retries):
        try:
            if provider == "openai":
                import openai
                openai.OpenAI(api_key=api_key).models.list()
            elif provider == "anthropic":
                import anthropic
                anthropic.Anthropic(api_key=api_key).messages.create(
                    model="claude-haiku-4-5-20251001", max_tokens=1,
                    messages=[{"role": "user", "content": "hi"}])
            elif provider == "azure":
                import openai
                openai.AzureOpenAI(
                    api_key=api_key, azure_endpoint=endpoint, api_version="2024-02-01"
                ).models.list()
            else:
                return jsonify({"connected": False, "error": f"Unknown provider: {provider}"})
            return jsonify({"connected": True})
        except Exception as e:
            last_err = str(e)[:200]
            # Retry on transient errors (overloaded, rate limit, 5xx)
            if any(code in last_err for code in ("529", "overloaded", "rate", "500", "502", "503")):
                if attempt < max_retries - 1:
                    _time.sleep(2 * (attempt + 1))
                    continue
            # Auth errors should not retry
            break
    return jsonify({"connected": False, "error": last_err})


@app.route("/api/settings/delete-key", methods=["POST"])
def delete_api_key():
    current = load_settings()
    current["llm_api_key"] = ""
    current["llm_provider"] = ""
    current["llm_endpoint"] = ""
    save_settings(current)
    return jsonify({"status": "ok"})


# ─── Rubric API ───

@app.route("/api/rubrics", methods=["GET"])
def get_rubrics():
    return jsonify(load_rubrics())


@app.route("/api/rubrics", methods=["POST"])
def add_rubric():
    data = request.get_json()
    rubrics = load_rubrics()
    rubric = {
        "id": str(uuid.uuid4())[:8],
        "name": data.get("name", "Untitled Rubric"),
        "criteria": data.get("criteria", []),
        "created": time.strftime("%Y-%m-%d %H:%M"),
    }
    rubrics.append(rubric)
    save_rubrics(rubrics)
    return jsonify({"status": "ok", "rubric": rubric})


@app.route("/api/rubrics/<rubric_id>", methods=["DELETE"])
def delete_rubric(rubric_id):
    rubrics = [r for r in load_rubrics() if r["id"] != rubric_id]
    save_rubrics(rubrics)
    return jsonify({"status": "ok"})


@app.route("/api/rubrics/<rubric_id>", methods=["PUT"])
def update_rubric(rubric_id):
    """Update an existing rubric."""
    data = request.get_json()
    rubrics = load_rubrics()
    for r in rubrics:
        if r["id"] == rubric_id:
            if "name" in data:
                r["name"] = data["name"]
            if "criteria" in data:
                r["criteria"] = data["criteria"]
            break
    save_rubrics(rubrics)
    return jsonify({"status": "ok"})


@app.route("/api/rubrics/suggest", methods=["POST"])
def suggest_criteria():
    """Use the LLM to suggest rubric criteria from a broad description."""
    data = request.get_json()
    description = data.get("description", "").strip()
    if not description:
        return jsonify({"error": "Please describe what you would like to assess."}), 400

    try:
        system = (
            "You are an education assessment expert specialising in Minecraft Education activities. "
            "Based on the user's description of what they want to assess, suggest appropriate rubric criteria. "
            "Return ONLY valid JSON with this structure: "
            '{"name": "Suggested Rubric Name", "criteria": [{"name": "Criterion Name", "description": "What to assess and what evidence to look for in Minecraft activity data"}]}'
            " Suggest 3-6 focused, observable criteria relevant to Minecraft Education gameplay."
        )
        result = _llm_chat(description, system=system)
        start = result.find("{")
        end = result.rfind("}") + 1
        if start >= 0 and end > start:
            suggestion = json.loads(result[start:end])
        else:
            raise ValueError("LLM did not return valid JSON")
        return jsonify({"status": "ok", "suggestion": suggestion})
    except Exception as e:
        return jsonify({"error": str(e)[:300]}), 500


@app.route("/api/rubrics/generate", methods=["POST"])
def generate_rubric():
    """Use the LLM to generate a rubric from uploaded file content."""
    file_text = ""
    if "file" in request.files:
        f = request.files["file"]
        filename = (f.filename or "").lower()
        raw_bytes = f.read()

        if filename.endswith(".pdf"):
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(stream=raw_bytes, filetype="pdf")
                pages = []
                for page in doc:
                    pages.append(page.get_text())
                doc.close()
                file_text = "\n\n".join(pages)
            except Exception as e:
                return jsonify({"error": f"Failed to read PDF: {e}"}), 400
        else:
            file_text = raw_bytes.decode("utf-8", errors="replace")

    elif request.is_json:
        file_text = request.get_json().get("text", "")

    if not file_text or not file_text.strip():
        return jsonify({"error": "No readable content found in the file"}), 400

    try:
        system = (
            "You are an education assessment expert. Generate a rubric from the provided content. "
            "Return ONLY valid JSON with this structure: "
            '{"name": "Rubric Name", "criteria": [{"name": "Criterion", "description": "What to assess", '
            '"levels": [{"label": "Excellent", "points": 4, "description": "..."}, '
            '{"label": "Good", "points": 3, "description": "..."}, '
            '{"label": "Developing", "points": 2, "description": "..."}, '
            '{"label": "Beginning", "points": 1, "description": "..."}]}]}'
        )
        result = _llm_chat(file_text[:8000], system=system)

        # Extract JSON from response
        start = result.find("{")
        end = result.rfind("}") + 1
        if start >= 0 and end > start:
            rubric_data = json.loads(result[start:end])
        else:
            raise ValueError("LLM did not return valid JSON")

        return jsonify({"status": "ok", "rubric": rubric_data})
    except Exception as e:
        return jsonify({"error": str(e)[:300]}), 500


# ─── Assessment API ───

@app.route("/api/player-stats", methods=["GET"])
def get_player_stats():
    """Return current player stats for the frontend."""
    stats = mc_server.get_player_stats()
    return jsonify(stats)


_plt = None

def _get_plt():
    global _plt
    if _plt is None:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        _plt = plt
    return _plt


def _make_movement_graph(player_name, trail):
    """Generate a 2D movement graph (XZ plane) and return as bytes."""
    plt = _get_plt()

    fig, ax = plt.subplots(1, 1, figsize=(5, 4))
    if not trail:
        ax.text(0.5, 0.5, "No movement data", ha="center", va="center",
                transform=ax.transAxes, color="#a19f9d")
    else:
        xs = [p[1] for p in trail]
        zs = [p[3] for p in trail]
        ax.plot(xs, zs, linewidth=1.2, color="#0078d4", alpha=0.7)
        ax.plot(xs[0], zs[0], "o", color="#107c10", markersize=8, label="Start")
        ax.plot(xs[-1], zs[-1], "s", color="#d13438", markersize=8, label="End")
        ax.legend(fontsize=8)

    ax.set_xlabel("X")
    ax.set_ylabel("Z")
    ax.set_title(f"{player_name} - Movement Path (Top-Down)", fontsize=11)
    ax.set_aspect("equal", adjustable="datalim")
    ax.grid(True, alpha=0.3)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=120)
    plt.close(fig)
    buf.seek(0)
    return buf


def _format_duration(seconds):
    m, s = divmod(int(seconds), 60)
    h, m = divmod(m, 60)
    if h:
        return f"{h}h {m}m {s}s"
    return f"{m}m {s}s"


def _build_player_summary(name, stats):
    """Build a compact data summary for one player."""
    s = stats.get(name, {})
    duration = s.get("time_connected_seconds", 0)
    # Top 5 block types only
    bp = dict(sorted(s.get("blocks_placed_types", {}).items(), key=lambda x: -x[1])[:5])
    bb = dict(sorted(s.get("blocks_broken_types", {}).items(), key=lambda x: -x[1])[:5])
    return {
        "time": _format_duration(duration),
        "distance": round(s.get("distance_travelled", 0), 1),
        "blocks_placed": s.get("blocks_placed", 0),
        "blocks_broken": s.get("blocks_broken", 0),
        "top_placed": bp,
        "top_broken": bb,
        "items_acquired": s.get("items_acquired", 0),
        "items_used": s.get("items_used", 0),
        "mobs_killed": s.get("mobs_killed", 0),
        "messages": s.get("messages_sent", 0),
        "chat": [m["text"] for m in s.get("messages", [])[-10:]],
        "events": s.get("events_total", 0),
    }


DETAIL_INSTRUCTIONS = {
    1: ("Give a single sentence per criterion. Synoptic assessment: 1-2 sentences.", 512),
    2: ("Give 1-2 sentences per criterion citing key numbers. Synoptic assessment: 2-3 sentences.", 768),
    3: ("Give a short paragraph per criterion with specific evidence from the data. Synoptic assessment: a paragraph.", 1024),
    4: ("Give a detailed paragraph per criterion analysing the data thoroughly, noting patterns and context. Synoptic assessment: a detailed paragraph considering all criteria.", 1536),
    5: ("Give a comprehensive, in-depth analysis per criterion covering all available evidence, patterns, comparisons, and implications. Synoptic assessment: a thorough multi-paragraph holistic analysis.", 2048),
}


def _assess_one_player(name, summary, criteria_text, rubric_name, detail_level=3):
    """Call the LLM to assess a single player. Returns parsed dict."""
    detail_instruction, max_tok = DETAIL_INSTRUCTIONS.get(detail_level, DETAIL_INSTRUCTIONS[3])

    prompt = (
        f"Assess Minecraft Education player \"{name}\" against this rubric.\n\n"
        f"RUBRIC: {rubric_name}\nCRITERIA:\n{criteria_text}\n\n"
        f"PLAYER DATA: {json.dumps(summary)}\n\n"
        f"INSTRUCTIONS: No grades/scores. {detail_instruction} "
        f"Set sufficient_data=false if data is lacking. "
        f"Cite specific numbers and actions from the data.\n\n"
        f"Return ONLY JSON: {{\"criteria_assessments\": [{{\"criterion\": \"...\", "
        f"\"observation\": \"...\", \"sufficient_data\": true/false}}], "
        f"\"synoptic_assessment\": \"...\"}}"
    )
    result = _llm_chat(prompt,
                        system="Education assessment expert. Return ONLY valid JSON. No grades.",
                        max_tokens=max_tok)
    start = result.find("{")
    end = result.rfind("}") + 1
    return json.loads(result[start:end])


@app.route("/api/assess", methods=["POST"])
def run_assessment():
    """Assess selected players against a rubric using the LLM, return Word doc."""
    data = request.get_json()
    rubric_id = data.get("rubric_id")
    player_names = data.get("players", [])

    rubrics = load_rubrics()
    rubric = next((r for r in rubrics if r["id"] == rubric_id), None)
    if not rubric:
        return jsonify({"error": "Rubric not found"}), 404

    all_stats = mc_server.get_player_stats()
    selected_stats = {n: all_stats[n] for n in player_names if n in all_stats}

    if not selected_stats:
        return jsonify({"error": "No player data available. Players must connect and generate activity before an assessment can be run."}), 400

    settings = load_settings()
    provider = settings.get("llm_provider", "unknown")
    api_key = settings.get("llm_api_key", "")
    if not provider or not api_key:
        return jsonify({"error": "No LLM API key configured. Add one in Settings."}), 400

    detail_level = settings.get("report_detail_level", 3)
    _, _, model_name = _get_llm_client()
    model_label = f"{provider.upper()} / {model_name}" if model_name else "Unknown"
    assessment_time = time.strftime("%Y-%m-%d %H:%M:%S")
    detail_labels = {1: "Minimal", 2: "Brief", 3: "Standard", 4: "Detailed", 5: "Comprehensive"}
    print(f"[ASSESS] Starting: {len(player_names)} players, rubric={rubric['name']}, detail={detail_level}", flush=True)

    criteria_text = "\n".join(
        f"- {c['name']}: {c['description']}" for c in rubric.get("criteria", [])
    )

    # Assess each player individually (faster, smaller prompts)
    assessments = []
    player_summaries = {}
    for idx, name in enumerate(player_names):
        if name not in selected_stats:
            continue
        print(f"[ASSESS] Player {idx+1}/{len(player_names)}: {name}", flush=True)
        summary = _build_player_summary(name, selected_stats)
        player_summaries[name] = summary
        try:
            result = _assess_one_player(name, summary, criteria_text, rubric["name"], detail_level)
            result["player"] = name
            assessments.append(result)
            print(f"[ASSESS] Player {name}: OK", flush=True)
        except Exception as e:
            print(f"[ASSESS] Player {name}: FAILED - {e}", flush=True)
            assessments.append({
                "player": name,
                "criteria_assessments": [{"criterion": "Error", "observation": str(e)[:200], "sufficient_data": False}],
                "synoptic_assessment": f"Assessment could not be completed: {str(e)[:100]}"
            })

    # ── Generate Word document ──
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Skills Crafter Assessment Report", level=0)
    doc.add_paragraph("")

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Light List Accent 1"
    for i, (label, val) in enumerate([
        ("Date & Time", assessment_time),
        ("AI Model", model_label),
        ("Report Detail", f"Level {detail_level} — {detail_labels.get(detail_level, 'Standard')}"),
        ("Rubric", rubric["name"]),
        ("Players Assessed", ", ".join(player_names)),
        ("Criteria Count", str(len(rubric.get("criteria", [])))),
    ]):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = val

    doc.add_paragraph("")
    doc.add_heading("Rubric Criteria", level=1)
    for c in rubric.get("criteria", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{c['name']}: ")
        run.bold = True
        p.add_run(c.get("description", ""))
    doc.add_paragraph("")

    for entry in assessments:
        pname = entry.get("player", "Unknown")
        ps = player_summaries.get(pname, {})
        ss = selected_stats.get(pname, {})

        doc.add_heading(f"Player: {pname}", level=1)

        # Activity summary
        doc.add_heading("Activity Summary", level=2)
        stat_table = doc.add_table(rows=1, cols=2)
        stat_table.style = "Light List Accent 1"
        stat_table.rows[0].cells[0].text = "Metric"
        stat_table.rows[0].cells[1].text = "Value"
        for label, val in [
            ("Time Connected", ps.get("time", "N/A")),
            ("Distance Travelled", f"{ps.get('distance', 0)} blocks"),
            ("Blocks Placed", str(ps.get("blocks_placed", 0))),
            ("Blocks Broken", str(ps.get("blocks_broken", 0))),
            ("Items Acquired", str(ps.get("items_acquired", 0))),
            ("Mobs Killed", str(ps.get("mobs_killed", 0))),
            ("Chat Messages", str(ps.get("messages", 0))),
            ("Total Events", str(ps.get("events", 0))),
        ]:
            row = stat_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = val

        bp = ps.get("top_placed", {})
        if bp:
            p = doc.add_paragraph()
            run = p.add_run("Top Blocks Placed: ")
            run.bold = True
            p.add_run(", ".join(f"{b} ({c})" for b, c in bp.items()))

        doc.add_paragraph("")

        # Movement graph
        doc.add_heading("Movement Path", level=2)
        trail = ss.get("position_trail", [])
        graph_buf = _make_movement_graph(pname, trail)
        doc.add_picture(graph_buf, width=Inches(4.5))
        doc.add_paragraph("")

        # Criterion assessments
        doc.add_heading("Criterion Assessments", level=2)
        for ca in entry.get("criteria_assessments", []):
            p = doc.add_paragraph()
            run = p.add_run(ca.get("criterion", "") + ": ")
            run.bold = True
            if not ca.get("sufficient_data", True):
                warn = p.add_run("[INSUFFICIENT DATA] ")
                warn.bold = True
                warn.font.color.rgb = RGBColor(0xD1, 0x34, 0x38)
            p.add_run(ca.get("observation", ""))

        doc.add_paragraph("")
        doc.add_heading("Synoptic Assessment", level=2)
        doc.add_paragraph(entry.get("synoptic_assessment", "No synoptic assessment available."))
        doc.add_page_break()

    p = doc.add_paragraph()
    run = p.add_run(f"Report generated by Skills Crafter on {assessment_time} using {model_label}. "
                     "This is an AI-assisted qualitative assessment based on observed in-game activity data. "
                     "No grades have been assigned.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x5E, 0x5C)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"assessment-{time.strftime('%Y%m%d-%H%M%S')}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ─── AI Chat System ───

# Per-player conversation history: {player_name: [{role, content}, ...]}
player_conversations = {}
chat_enabled = False

AI_CHAT_SYSTEM = (
    "You are a helpful Minecraft Education assistant called Skills Crafter AI. "
    "Keep responses short (1-3 sentences) and focused on Minecraft. "
    "Be friendly, encouraging, and educational. "
    "You can help with crafting recipes, game mechanics, building tips, and redstone. "
    "Do not discuss topics unrelated to Minecraft."
)

from build_toolkit import generate_build_commands, get_llm_build_prompt, STRUCTURE_SPECS


def _handle_chat_request(player_name, message, player_pos):
    """Process an @ai or @build chat message and send response back to Minecraft."""
    import threading

    def process():
        try:
            if message.lower().startswith("@build"):
                _handle_build(player_name, message[6:].strip(), player_pos)
            elif message.lower().startswith("@ai"):
                _handle_ai_chat(player_name, message[3:].strip())
        except Exception as e:
            print(f"[CHAT] Error for {player_name}: {e}", flush=True)
            _send_chat_response(player_name, f"Sorry, I encountered an error: {str(e)[:80]}")

    # Run in a thread so it doesn't block the queue reader
    threading.Thread(target=process, daemon=True).start()


def _handle_ai_chat(player_name, user_message):
    """Handle @ai conversation."""
    if not user_message:
        _send_chat_response(player_name, "Hi! Ask me anything about Minecraft. Use @ai followed by your question.")
        return

    # Get or create conversation history
    if player_name not in player_conversations:
        player_conversations[player_name] = []

    history = player_conversations[player_name]
    history.append({"role": "user", "content": user_message})

    # Keep history manageable (last 10 exchanges)
    if len(history) > 20:
        history[:] = history[-20:]

    print(f"[CHAT] @ai from {player_name}: {user_message}", flush=True)

    client, provider, model = _get_llm_client()
    if not client:
        _send_chat_response(player_name, "AI is not configured. Ask the server operator to add an API key.")
        return

    try:
        if provider == "anthropic":
            resp = client.messages.create(
                model=model, max_tokens=150, system=AI_CHAT_SYSTEM,
                messages=history, timeout=20)
            reply = resp.content[0].text
        else:
            messages = [{"role": "system", "content": AI_CHAT_SYSTEM}] + history
            resp = client.chat.completions.create(
                model=model, messages=messages, max_tokens=150, timeout=20)
            reply = resp.choices[0].message.content

        history.append({"role": "assistant", "content": reply})
        print(f"[CHAT] Reply to {player_name}: {reply[:80]}", flush=True)
        _send_chat_response(player_name, reply)
    except Exception as e:
        print(f"[CHAT] LLM error: {e}", flush=True)
        _send_chat_response(player_name, "Sorry, I couldn't process that right now. Try again in a moment.")


def _handle_build(player_name, description, pos):
    """Handle @build command — LLM generates a spec, toolkit builds it."""
    if not description:
        types = ", ".join(STRUCTURE_SPECS.keys())
        _send_chat_response(player_name, f"Tell me what to build! Types: {types}. Example: @build large stone house")
        return

    print(f"[BUILD] @build from {player_name}: {description}", flush=True)

    client, provider, model = _get_llm_client()
    if not client:
        _send_chat_response(player_name, "AI is not configured.")
        return

    _send_chat_response(player_name, f"Planning your {description}...")

    try:
        # Step 1: LLM generates a structured build spec
        system = get_llm_build_prompt()
        prompt = f"Build request: {description}\nReturn ONLY the JSON spec object for this specific build. Do not repeat previous builds."
        print(f"[BUILD] Sending to LLM: {prompt}", flush=True)
        result = _llm_chat(prompt, system=system, max_tokens=300, timeout=20)
        print(f"[BUILD] LLM raw response: {result[:300]}", flush=True)

        # Extract JSON spec
        start = result.find("{")
        end = result.rfind("}") + 1
        if start < 0 or end <= start:
            _send_chat_response(player_name, "Sorry, I couldn't plan that build. Try something like: @build house")
            return

        spec = json.loads(result[start:end])
        struct_type = spec.get("type", "unknown")
        print(f"[BUILD] LLM spec: type={struct_type}, {json.dumps(spec)}", flush=True)

        # Step 2: Toolkit generates precise commands
        commands = generate_build_commands(spec)

        _send_chat_response(player_name, f"Building {struct_type} ({len(commands)} commands)...")

        # Step 3: Execute commands with small batches
        for cmd in commands:
            mc_server.send_command(cmd)

        _send_chat_response(player_name, f"Done! Built a {struct_type} with {len(commands)} blocks.")
        print(f"[BUILD] Executed {len(commands)} commands for {player_name}", flush=True)

    except Exception as e:
        print(f"[BUILD] Error: {e}", flush=True)
        _send_chat_response(player_name, "Sorry, the build failed. Try: @build house, @build tower, @build fountain")


def _send_chat_response(player_name, message):
    """Send a chat message back to a specific player in Minecraft."""
    # Use tellraw for formatted response
    raw = json.dumps({"rawtext": [{"text": "\u00a7b[AI] \u00a7f" + message}]})
    mc_server.send_command(f'tellraw "{player_name}" {raw}')


mc_server.chat_handler = _handle_chat_request


# ─── Game Commands API ───

@socketio.on("game_command")
def handle_game_command(data):
    """Send a Minecraft command to connected players."""
    command = data.get("command", "")
    targets = data.get("targets", "@a")  # default all players

    if not mc_server.running:
        socketio.emit("command_result", {"success": False, "error": "Server not running"})
        return

    # Replace @s with specific targets if needed
    final_cmd = command.replace("@s", targets)
    mc_server.send_command(final_cmd)
    socketio.emit("command_result", {"success": True, "command": final_cmd})


# ─── WebSocket Events ───

@socketio.on("set_chat")
def handle_set_chat(data):
    global chat_enabled
    chat_enabled = data.get("enabled", False)
    mc_server.chat_enabled = chat_enabled
    # Update the flag in the subprocess
    if mc_server._chat_enabled_flag:
        if chat_enabled:
            mc_server._chat_enabled_flag.set()
        else:
            mc_server._chat_enabled_flag.clear()
    print(f"[CHAT] AI chat {'enabled' if chat_enabled else 'disabled'}", flush=True)


@socketio.on("start_server")
def handle_start_server():
    if not mc_server.running:
        settings = load_settings()
        mc_server.welcome_message = settings["welcome_message"]
        mc_server.welcome_color = settings["welcome_color"]
        mc_server.chat_enabled = chat_enabled
        mc_server.start()


@socketio.on("stop_server")
def handle_stop_server():
    if mc_server.running:
        mc_server.stop()


@socketio.on("connect")
def handle_web_connect():
    socketio.emit("server_status", {
        "running": mc_server.running,
        "host": mc_server.host,
        "port": mc_server.port,
    })
    socketio.emit("players_update", mc_server._get_players_data())


if __name__ == "__main__":
    settings = load_settings()
    mc_server.welcome_message = settings["welcome_message"]
    mc_server.welcome_color = settings["welcome_color"]

    print("=" * 50)
    print("  Skills Crafter - Minecraft Education Dashboard")
    print("=" * 50)
    print(f"  Open http://localhost:5050 in your browser")
    print(f"  Local IP: {get_public_host()}")
    print("=" * 50)
    socketio.run(app, host="0.0.0.0", port=5050, debug=False, allow_unsafe_werkzeug=True)
